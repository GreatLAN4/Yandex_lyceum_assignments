from docx import Document
from sys import stdin as sys

document = Document()

document.add_paragraph(input()).bold = True
document.add_paragraph(input())

document.add_paragraph("На празднования 8 марта приглашаются")
names = document.add_paragraph()
for name in sys:
    names.add_run(name).bold = True

document.save('Приглашения.docx')

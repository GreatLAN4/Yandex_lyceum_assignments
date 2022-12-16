from docx import Document

document = Document()


def markdown_to_docx(text):
    linesINtext = text.split('\n')
    file_name = linesINtext[0].strip()
    document.add_heading(file_name.strip(), 0)
    linesINtext.pop(0)
    flag_run = False
    viewer_flag_run = False
    list_flag = False

    for line in linesINtext:
        if flag_run:
            viewer_flag_run = True
            flag_run = False

        if line[0] != "#":

            if line[0] in '*-+' and line[1] == ' ':
                list_flag = True
                document.add_paragraph(line.strip("*-+ "), style='List Bullet')

            elif line[0].isdigit():
                list_flag = False
                if line.split('.')[0].isdigit():
                    document.add_paragraph(line.rstrip().lstrip("1234567890. "), style='List Number')
                else:
                    document.add_paragraph(line.strip())

            if line.startswith("*") or line.startswith("_"):
                n = list(line[:3])
                n = n.count(n[0])
                p = document.add_paragraph()
                run = p.add_run(line.strip("_* "))
                if n in [1, 3]:
                    run.italic = True
                if n in [2, 3]:
                    run.bold = True
            else:
                if not viewer_flag_run and not list_flag:
                    p = document.add_paragraph(line)
                else:
                    p.add_run(line)
                    viewer_flag_run = False

            if line[-2:] == "  ":
                flag_run = True

            if viewer_flag_run:
                p.add_run(line)
                viewer_flag_run, flag_run = False, False

        if line[0] == "#":
            a = list(line[:5]).count("#")
            line = line.strip("#").strip()
            document.add_heading(line, a)

    document.save('test.docx')


markdown_to_docx("test01\n"
                 "Абзацы создаются при помощи пустой строки. Если вокруг текста сверху и снизу есть пустые "
                 "строки,  то текст превращается в абзац.\n"
                 "Чтобы сделать перенос строки вместо абзаца,  \n"
                 "нужно поставить два пробела в конце предыдущей строки.\n"
                 "Заголовки отмечаются диезом `#` в начале "
                 "строки, от одного до шести. Например:\n"
                 "# Заголовок первого уровня\n"
                 "## Заголовок h2\n"
                 "### Заголовок h3\n"
                 "#### Заголовок h4\n"
                 "##### Заголовок h5\n"
                 "###### Заголовок h6\n"
                 "В декоративных целях заголовки можно «закрывать» с обратной стороны.\n"
                 "### Списки\n"
                 "Для разметки неупорядоченных списков можно использовать или `*`, или `-,\n"
                 "- элемент 1\n"
                 "- элемент 2\n"
                 "- элемент ...\n"
                 "Вложенные пункты создаются четырьмя пробелами перед маркером пункта:\n"
                 "* элемент 1v"
                 "* элемент 2\n"
                 "    * вложенный элемент 2.1\n"
                 "    * вложенный элемент 2.2\n"
                 "* элемент ...\n"
                 "Упорядоченный список:\n"
                 "1. элемент 1\n"
                 "2. элемент 2\n"
                 "    1. вложенный\n"
                 "    2. вложенный\n"
                 "3. элемент 3\n"
                 "4. Donec sit amet nisl. Aliquam semper ipsum sit amet velit. Suspendisse id sem consectetuer libero "
                 "luctus adipiscing.\n "
                 "На самом деле не важно как в коде пронумерованы пункты, главное, чтобы перед элементом списка "
                 "стояла цифра (любая) с точкой. Можно сделать и так:\n "
                 "0. элемент 1v"
                 "0. элемент 2\n"
                 "0. элемент 3\n"
                 "0. элемент 4"
                 )
